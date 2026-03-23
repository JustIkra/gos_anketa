"""Parse a DOCX questionnaire file into structured institution data."""

import re
from dataclasses import dataclass, field

from docx import Document


@dataclass
class ParsedInstitution:
    full_name: str = ""
    short_name: str = ""
    contacts: dict = field(default_factory=lambda: {
        "address": "", "phone": "", "fax": "", "email": "", "website": "",
    })
    personnel: list[dict] = field(default_factory=list)
    programs: list[dict] = field(default_factory=list)


_CODE_PATTERN = re.compile(r"^\d{2}\.\d{2}\.\d{2}$")

_CONTACT_KEYS_MAP = {
    "адрес": "address",
    "телефон": "phone",
    "факс": "fax",
    "e-mail": "email",
    "email": "email",
    "сайт": "website",
}


def _cell_text(cell) -> str:
    """Return stripped text from a table cell, normalizing whitespace."""
    text = cell.text.strip() if cell.text else ""
    # Replace internal newlines with spaces
    return " ".join(text.split())


def _bool_value(text: str) -> bool:
    """Interpret table cell value as boolean for program forms."""
    t = text.strip()
    if not t or t in ("-", "---", "--", "—", "——"):
        return False
    return True


def _parse_contacts_table(table) -> dict:
    """Parse the contacts table (expected 5 rows x 2 cols)."""
    contacts = {"address": "", "phone": "", "fax": "", "email": "", "website": ""}
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = _cell_text(cells[0]).lower().rstrip(":")
        value = _cell_text(cells[1])
        for keyword, key in _CONTACT_KEYS_MAP.items():
            if keyword in label:
                contacts[key] = value
                break
    return contacts


def _parse_personnel_table(table) -> list[dict]:
    """Parse the personnel table. Row 0 is header, rest are data."""
    personnel = []
    for idx, row in enumerate(table.rows):
        if idx == 0:
            continue  # skip header
        cells = row.cells
        if len(cells) < 4:
            continue
        position = _cell_text(cells[0])
        full_name = _cell_text(cells[1])
        if not position and not full_name:
            continue
        work_phone = _cell_text(cells[2])
        mobile_phone = _cell_text(cells[3])
        personnel.append({
            "position": position,
            "full_name": full_name,
            "work_phone": work_phone,
            "mobile_phone": mobile_phone,
        })
    return personnel


def _parse_programs_table(table) -> list[dict]:
    """Parse the programs table with specialty/profession sections."""
    programs = []
    program_type = "specialty"
    sort_order = 0

    for row in table.rows:
        cells = row.cells
        if not cells:
            continue

        first_cell = _cell_text(cells[0])

        # Skip header rows with "Код"
        if first_cell.lower().startswith("код"):
            continue

        # Check for section header indicating switch to professions
        # These rows typically span all columns with the same text
        all_same = all(_cell_text(c) == first_cell for c in cells)
        if all_same and first_cell:
            lower = first_cell.lower()
            if "профессии" in lower or "рабочих" in lower or "служащих" in lower:
                program_type = "profession"
                continue
            if "специальност" in lower:
                program_type = "specialty"
                continue

        # Check if first cell looks like a program code
        code = first_cell.strip()
        if not _CODE_PATTERN.match(code):
            # Not a data row -- could be a section header
            if first_cell:
                lower = first_cell.lower()
                if "профессии" in lower or "рабочих" in lower or "служащих" in lower:
                    program_type = "profession"
                elif "специальност" in lower:
                    program_type = "specialty"
            continue

        # Data row
        specialty_name = _cell_text(cells[1]) if len(cells) > 1 else ""

        fulltime = False
        parttime = False
        distance = False

        if len(cells) > 2:
            fulltime = _bool_value(_cell_text(cells[2]))
        if len(cells) > 3:
            parttime = _bool_value(_cell_text(cells[3]))
        if len(cells) > 4:
            distance = _bool_value(_cell_text(cells[4]))

        programs.append({
            "code": code,
            "specialty_name": specialty_name,
            "program_type": program_type,
            "fulltime": fulltime,
            "parttime": parttime,
            "distance": distance,
            "sort_order": sort_order,
        })
        sort_order += 1

    return programs


def parse_docx(file_path: str) -> ParsedInstitution:
    """Parse a DOCX questionnaire and return structured data."""
    doc = Document(file_path)
    result = ParsedInstitution()

    # ── Extract names from paragraphs ────────────────────────────────────
    # Look for non-empty paragraphs before the first table.
    # Typically P2 (index 2) = full name, P4 (index 4) = short name.
    non_empty_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            non_empty_paras.append(text)

    # Full name is usually the first substantial paragraph,
    # short name comes a couple after.
    if len(non_empty_paras) >= 2:
        result.full_name = non_empty_paras[0]
        result.short_name = non_empty_paras[1]
    elif len(non_empty_paras) == 1:
        result.full_name = non_empty_paras[0]
        result.short_name = non_empty_paras[0]

    # Also try the positional approach: P2 and P4
    paragraphs = doc.paragraphs
    if len(paragraphs) > 4:
        p2_text = paragraphs[2].text.strip()
        p4_text = paragraphs[4].text.strip()
        if p2_text:
            result.full_name = p2_text
        if p4_text:
            result.short_name = p4_text

    # ── Parse tables ─────────────────────────────────────────────────────
    tables = doc.tables

    if len(tables) >= 1:
        result.contacts = _parse_contacts_table(tables[0])
    if len(tables) >= 2:
        result.personnel = _parse_personnel_table(tables[1])
    if len(tables) >= 3:
        result.programs = _parse_programs_table(tables[2])

    return result
